from docx import Document
from io import BytesIO

def generate_docx(data: dict) -> BytesIO:
    doc = Document()

    # Title
    doc.add_heading('Relazione Finale di Tirocinio', 0)

    # Setup Data
    setup = data.get('setup', {})
    doc.add_heading('Dati Anagrafici', level=1)
    doc.add_paragraph(f"Tirocinante: {setup.get('nomeTirocinante', 'N/A')}")
    doc.add_paragraph(f"Università/Percorso: {setup.get('universita', 'N/A')}")
    doc.add_paragraph(f"Scuola: {setup.get('scuola', 'N/A')}")
    doc.add_paragraph(f"Tutor: {setup.get('tutor', 'N/A')}")
    doc.add_paragraph(f"Classe: {setup.get('classe', 'N/A')}")

    # Generated Text
    gen_text = data.get('generatedText', {})

    doc.add_heading('Capitolo 1: Contesto Scuola', level=1)
    doc.add_paragraph(gen_text.get('capitolo1', 'Nessun testo generato.'))

    doc.add_heading('Capitolo 2: Contesto Classe', level=1)
    doc.add_paragraph(gen_text.get('capitolo2', 'Nessun testo generato.'))

    doc.add_heading('Capitolo 3: Attività Svolte', level=1)
    doc.add_paragraph(gen_text.get('capitolo3', 'Nessun testo generato.'))

    doc.add_heading('Conclusioni', level=1)
    doc.add_paragraph(gen_text.get('conclusioni', 'Nessun testo generato.'))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
